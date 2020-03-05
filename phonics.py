from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.api import Document

"""


2x4_phonics_slips.docx


"""
def set_cell_border( _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = _Cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


# read word list
def createdocs(docname,content):
 items = list(content)
 i = len(items)
 print(items, i)
 document = Document()

 #### set font
 style = document.styles['Normal']
 font = style.font
 font.name = 'Sofia Pro Soft'
 font.size = Pt(120)

 #### set Document/Margins
 for sec in document.sections:
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)

 ##### Table setting
 if (i % 2) == 0:
  nr=i/2
  table = document.add_table(rows=nr, cols=2, style='Table Grid')
 else:
     f=i+1
     nr = f / 2
     table = document.add_table(rows=nr, cols=2, style='Table Grid')
 for col in table.columns:
    for cell in col.cells:
        for par in cell.paragraphs:
            par.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
 d = 0  # type: int
 while d < nr:
     set_cell_border(
         table.cell(d, 0),
         top={"val": "dashed"},
         bottom={"val": "dashed"},
         start={"val": "dashed"},
         end={"val": "dashed"},
     )
     set_cell_border(
         table.cell(d, 1),
         top={"val": "dashed"},
         bottom={"val": "dashed"},
         start={"val": "dashed"},
         end={"val": "dashed"},
     )

     d = d + 1

 e = 0
 while e < nr:
    table.cell(e, 0).height = Cm(5.92)
    table.cell(e, 0).width = Cm(8.26)
    table.cell(e, 1).height = Cm(5.92)
    table.cell(e, 1).width = Cm(8.26)
    e = e + 1

 table.autofit = False
 #### Add table content
 if (i % 2) == 0:
  a = 0
  b = 0
  while b < nr:
    table.cell(b, 0).text = items[a]
    table.cell(b, 1).text = items[a+1]
    table.cell(b, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(b, 0).paragraphs[0].paragraph_format.space_before = Pt(18)
    table.cell(b, 1).paragraphs[0].paragraph_format.space_before = Pt(18)
    b = b + 1
    a = a + 2
  document.save(docname+'.docx')
 else:
     items.append(' ')
     i=len(items)
     nr=i/2
     a = 0
     b = 0
     while b < nr:
         print(b,a,items[a+1])
         table.cell(b, 0).text = items[a]
         table.cell(b, 1).text = items[a + 1]
         table.cell(b, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
         table.cell(b, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
         table.cell(b, 0).paragraphs[0].paragraph_format.space_before = Pt(18)
         table.cell(b, 1).paragraphs[0].paragraph_format.space_before = Pt(18)
         b = b + 1
         a = a + 2
     document.save(docname + '.docx')


document = Document('phonics_letter_slips_table.docx')
table = document.tables[0]
##Getting the original data from the document to a list
ls =[]
for column in table.columns:
    for cell in column.cells:
        for paragraph in cell.paragraphs:
            ls.append(paragraph.text)
ls = filter(None, ls)
i=len(ls)
print(ls)
print(i)

n=0
m=i/2
while n< i/2:
    docname=ls[n]
    content=ls[m]
    createdocs(docname, content)
    m=m+1
    n=n+1

